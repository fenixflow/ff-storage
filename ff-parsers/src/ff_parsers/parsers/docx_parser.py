"""
Word document parser implementation using python-docx.
"""

from pathlib import Path
from typing import Optional, Union, List
import docx
from docx.document import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from ..base import BaseParser, ParseOptions
from ..models import (
    ExtractedDocument,
    ExtractedText,
    ExtractedTable,
    ExtractedImage,
    DocumentMetadata,
)
from ..exceptions import CorruptedFileError
from ..utils.cleaning import clean_text, remove_extra_spaces_in_table


class DocxParser(BaseParser):
    """Parser for Word documents (.docx)."""

    def get_supported_extensions(self) -> list[str]:
        """Get supported file extensions."""
        return [".docx", ".DOCX"]

    def get_mime_types(self) -> list[str]:
        """Get supported MIME types."""
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]

    def validate(self, file_path: Union[str, Path]) -> bool:
        """Validate that the file is a valid Word document."""
        try:
            path = Path(file_path)
            if not path.exists() or not path.is_file():
                return False

            # Check extension
            if path.suffix.lower() != ".docx":
                return False

            # Try to open with python-docx
            doc = docx.Document(path)
            # Check if we can access paragraphs
            _ = doc.paragraphs
            return True

        except Exception:
            return False

    def extract_metadata(self, file_path: Union[str, Path]) -> DocumentMetadata:
        """Extract Word document metadata."""
        path = self._ensure_path(file_path)
        metadata = DocumentMetadata()

        try:
            metadata.file_size = self._get_file_size(path)
            metadata.mime_type = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # Load document
            doc = docx.Document(path)

            # Access core properties
            props = doc.core_properties
            if props:
                metadata.title = props.title
                metadata.author = props.author
                metadata.subject = props.subject
                metadata.created_date = props.created
                metadata.modified_date = props.modified

                if props.keywords:
                    metadata.keywords = [k.strip() for k in props.keywords.split(",")]

                # Custom properties
                metadata.custom_properties["last_modified_by"] = props.last_modified_by
                metadata.custom_properties["revision"] = props.revision
                metadata.custom_properties["category"] = props.category
                metadata.custom_properties["comments"] = props.comments

            # Count pages (approximate - Word doesn't store exact page count)
            # We'll count based on paragraphs as an estimate
            metadata.page_count = self._estimate_page_count(doc)

        except Exception as e:
            self._log_warning(f"Failed to extract metadata: {e}")

        return metadata

    def parse(
        self, file_path: Union[str, Path], options: Optional[ParseOptions] = None
    ) -> ExtractedDocument:
        """Parse Word document and extract content."""
        path = self._ensure_path(file_path)
        options = options or self._get_default_options()

        # Create base document
        document = self._create_base_document(path)

        try:
            # Load Word document
            doc = docx.Document(path)

            # Extract metadata
            if options.extract_metadata:
                document.metadata = self.extract_metadata(path)

            # Extract content
            all_text = []
            paragraph_texts = []
            current_page = 1
            paragraph_count = 0

            for element in self._iter_block_items(doc):
                if isinstance(element, Paragraph):
                    # Extract paragraph text
                    para_text = self._extract_paragraph_text(element, options)
                    if para_text:
                        paragraph_count += 1
                        paragraph_texts.append(para_text)
                        all_text.append(para_text)

                        # Create ExtractedText for this paragraph
                        extracted = ExtractedText(
                            content=para_text,
                            page=current_page,
                            paragraph=paragraph_count,
                            confidence=1.0,
                        )
                        document.pages.append(extracted)

                        # Estimate page breaks (rough approximation)
                        if paragraph_count % 30 == 0:  # Approximately 30 paragraphs per page
                            current_page += 1

                elif isinstance(element, Table) and options.extract_tables:
                    # Extract table
                    table = self._extract_table(element, current_page)
                    if table:
                        document.tables.append(table)
                        # Add table text representation
                        table_text = self._table_to_text(table)
                        all_text.append(table_text)

            # Combine all text
            document.text = "\n\n".join(all_text)

            # Extract images if requested
            if options.extract_images:
                images = self._extract_images(doc)
                document.images.extend(images)

            # Count words
            if document.text:
                document.metadata.word_count = len(document.text.split())

        except Exception as e:
            raise CorruptedFileError(
                f"Failed to parse Word document: {str(e)}", partial_result=document
            )

        return document

    def _iter_block_items(self, document: Document):
        """
        Yield each paragraph and table in the document in order.
        """
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        parent = document.element.body
        for child in parent.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def _extract_paragraph_text(self, paragraph: Paragraph, options: ParseOptions) -> Optional[str]:
        """Extract text from a paragraph."""
        try:
            text = paragraph.text

            if not text or not text.strip():
                return None

            # Clean text if requested
            if not options.preserve_whitespace:
                text = clean_text(text, preserve_paragraphs=False)

            return text

        except Exception as e:
            self._log_warning(f"Failed to extract paragraph text: {e}")
            return None

    def _extract_table(self, table: Table, page_num: int) -> Optional[ExtractedTable]:
        """Extract data from a Word table."""
        try:
            if not table.rows:
                return None

            # Extract headers (first row)
            headers = []
            first_row = table.rows[0]
            for cell in first_row.cells:
                headers.append(cell.text.strip())

            # Extract data rows
            rows = []
            for row in table.rows[1:]:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                rows.append(row_data)

            # Clean table data
            rows = remove_extra_spaces_in_table(rows)

            # Skip empty tables
            if not headers or not rows:
                return None

            return ExtractedTable(
                headers=headers, rows=rows, page=page_num, name=f"Table_{page_num}_{len(rows)}"
            )

        except Exception as e:
            self._log_warning(f"Failed to extract table: {e}")
            return None

    def _table_to_text(self, table: ExtractedTable) -> str:
        """Convert table to text representation."""
        lines = []

        # Add headers
        lines.append("\t".join(table.headers))

        # Add rows
        for row in table.rows:
            lines.append("\t".join(row))

        return "\n".join(lines)

    def _extract_images(self, doc: Document) -> List[ExtractedImage]:
        """Extract image metadata from document."""
        images = []

        try:
            # Access document relationships
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # Get image data
                        image_part = rel.target_part
                        # image_data = image_part.blob  # Available if needed in future

                        # Create image metadata
                        image = ExtractedImage(
                            format=image_part.content_type.split("/")[-1].upper(),
                            image_data=None,  # Don't store actual data to save memory
                        )
                        images.append(image)

                    except Exception as e:
                        self._log_warning(f"Failed to extract image: {e}")

        except Exception as e:
            self._log_warning(f"Failed to extract images: {e}")

        return images

    def _estimate_page_count(self, doc: Document) -> int:
        """Estimate page count based on content."""
        # Rough estimation: 30 paragraphs per page
        paragraph_count = len(doc.paragraphs)
        table_count = len(doc.tables)

        # Each table counts as roughly 5 paragraphs
        estimated_items = paragraph_count + (table_count * 5)

        # Estimate pages
        pages = max(1, (estimated_items + 29) // 30)

        return pages
